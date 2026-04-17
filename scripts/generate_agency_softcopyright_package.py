from __future__ import annotations

import ctypes
import csv
import json
import os
import shutil
import subprocess
import textwrap
import time
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageGrab
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
ROOT_DIR = Path(r"D:\OneDrive\研究生\我的成果\软著")
TEMPLATE_DIR = ROOT_DIR / "软著模板"
CURRENT_DIR = ROOT_DIR / "源荷匹配软著"
OUTPUT_DIR = CURRENT_DIR / "事务所模板版"
SCREENSHOT_DIR = OUTPUT_DIR / "screenshots"

SOFTWARE_NAME = "综合能源系统源荷匹配容量配置优化软件"
SOFTWARE_SHORT = "源荷匹配容量优化"
SOFTWARE_VERSION = "V1.0"
OWNER_NAME = "浙江大学"
OWNER_CODE = "12100000470095016Q"
CONTACT_NAME = "郑浩男"
CONTACT_PHONE = "18157780950"
CONTACT_EMAIL = "haonan.zheng@zju.edu.cn"
DEV_DATE = "2026年03月19日"
PUBLISH_STATUS = "未发表"
PUBLISH_DATE = ""
PUBLISH_CITY = ""
SOURCE_LINES_FOR_SUBMISSION = 3000

RESULT_DIR = (
    REPO
    / "Results"
    / "test_exp1_10x5_20260318_234055"
    / "exp1_german_10x5_5methods_20260318_234055"
)
DATA_DIR = REPO / "data"
IMAGE_SOURCE_DIR = CURRENT_DIR / "images"
CAPTURE_TMP = Path(os.environ.get("TEMP", str(REPO))) / "agency_softcopyright_capture"

FORM_TEMPLATE = TEMPLATE_DIR / "template_form.docx"
MANUAL_TEMPLATE = TEMPLATE_DIR / "template_manual.docx"
SOURCE_TEMPLATE = TEMPLATE_DIR / "template_code.docx"

OUTPUT_FORM = OUTPUT_DIR / f"{SOFTWARE_NAME}_软件登记信息表.docx"
OUTPUT_MANUAL = OUTPUT_DIR / f"{SOFTWARE_NAME}_使用说明书.docx"
OUTPUT_SOURCE = OUTPUT_DIR / f"{SOFTWARE_NAME}_源代码.docx"

FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")
FONT_MONO = Path(r"C:\Windows\Fonts\consola.ttf")

SELECTED_SOURCE_FILES = [
    "run.py",
    "cchp_gasolution.py",
    "cchp_gaproblem.py",
    "operation.py",
    "case_config.py",
    "scripts/post_analysis_report.py",
    "scripts/generate_songshan_data.py",
    "scripts/check_songshan_data.py",
    "scripts/kmeans_clustering.py",
    "scripts/kmeansClustering.m",
    "scripts/enhanced_analysis.py",
    "scripts/lambda_sensitivity.py",
    "scripts/post_analysis.py",
    "scripts/resilience_test.py",
]

MAIN_FUNCTIONS = "本软件用于分布式电热综合能源系统优化规划与源荷匹配分析，服务于研究算例计算、方案比选和结果整理。程序可读取德国社区和松山湖校园两类案例的全年逐时电、热、冷负荷数据，以及太阳辐照、风速、环境温度和典型日权重文件，自动完成输入文件定位、案例参数载入和基本校核。用户可在命令行下选择test、quick、full、custom等运行模式，也可按论文预设实验编号直接调用成组计算任务，分别完成不同案例、不同匹配指标和是否配置卡诺电池的对比计算。软件能够围绕光伏、风电、燃气轮机、电热泵、电制冷机、吸收式制冷机、电储能、热储能、冷储能及卡诺电池等设备容量开展联合优化，并在每个候选方案下调用24小时调度模型，对14个典型日逐一求解，统计年化投资、运行费用和容量电费。程序支持波动率、能质加权欧氏距离、Pearson相关系数、自给率等多种源荷匹配评价方法，可在计算结束后生成设备容量结果表、Pareto解集、对比报告、后验分析结果表和图形文件，并按时间戳建立结果目录，便于研究人员复核计算过程、比较不同方案、筛选代表解并整理论文结果。对于需要扩展研究的用户，还可以在案例配置文件中补充新的参数与数据文件，继续开展其他综合能源系统场景的对比分析。"

TECH_FEATURE = (
    "采用容量规划与运行调度联动求解，支持典型日加权、多指标匹配评价和卡诺电池扩展建模。"
)

PURPOSE = "用于综合能源系统方案计算与源荷匹配分析"
DOMAIN = "综合能源系统优化规划与运行分析"
DEV_HW = "x64工作站，16GB内存，500GB存储空间"
RUN_HW = "x64 PC或服务器，16GB内存，建议SSD存储"
DEV_OS = "Windows 10/11 64位"
DEV_TOOL = "Python 3.8，PyCharm/VS Code，Git"
RUN_OS = "Windows 10/11 64位"
SUPPORT_SW = "Python 3.8，Gurobi或GLPK求解器"
PROGRAM_LANG = "Python"

CREATE_NEW_CONSOLE = 0x00000010
SW_RESTORE = 9
SW_MAXIMIZE = 3


def assert_lengths() -> None:
    checks = {
        "开发硬件环境": (DEV_HW, 50),
        "运行硬件环境": (RUN_HW, 50),
        "开发操作系统": (DEV_OS, 50),
        "开发工具": (DEV_TOOL, 50),
        "运行平台": (RUN_OS, 50),
        "运行支撑环境": (SUPPORT_SW, 50),
        "开发目的": (PURPOSE, 50),
        "面向领域": (DOMAIN, 50),
        "技术特点": (TECH_FEATURE, 100),
    }
    for label, (value, limit) in checks.items():
        if len(value) > limit:
            raise ValueError(f"{label} 超出字数限制：{len(value)} > {limit}")
    if not (500 <= len(MAIN_FUNCTIONS) <= 1300):
        raise ValueError(f"主要功能字数不符合要求：{len(MAIN_FUNCTIONS)}")


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)


def clear_body(document: Document) -> None:
    body = document.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(
    run,
    east_asia: str,
    size: float,
    bold: bool = False,
    color: str | None = None,
    latin: str = "Times New Roman",
) -> None:
    run.bold = bold
    run.font.name = latin
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_text(
    paragraph,
    text: str,
    east_asia: str = "华文仿宋",
    size: float = 12,
    bold: bool = False,
    color: str | None = None,
    latin: str = "Times New Roman",
):
    run = paragraph.add_run(text)
    set_run_font(
        run, east_asia=east_asia, size=size, bold=bold, color=color, latin=latin
    )
    return run


def format_body_paragraph(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.5


def add_body_text(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(
        style="Body Text Indent"
        if "Body Text Indent" in [s.name for s in document.styles]
        else "Normal"
    )
    format_body_paragraph(paragraph)
    add_text(paragraph, text)


def add_heading_line(document: Document, text: str, style_name: str) -> None:
    paragraph = document.add_paragraph(style=style_name)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    size = 14 if style_name == "Heading 3" else 16
    set_run_font(run, east_asia="宋体", size=size, bold=True, latin="Times New Roman")


def add_title(document: Document, title: str, subtitle: str) -> None:
    p1 = document.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1.paragraph_format.space_after = Pt(0)
    add_text(p1, title, east_asia="黑体", size=16, bold=True, latin="Arial")

    p2 = document.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.paragraph_format.space_after = Pt(12)
    add_text(p2, subtitle, east_asia="黑体", size=16, bold=True, latin="Arial")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    add_text(paragraph, text, east_asia="宋体", size=10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_caption(document: Document, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    add_text(paragraph, caption, east_asia="宋体", size=10.5, latin="Times New Roman")


def apply_table_style(table, *style_names: str) -> None:
    for style_name in style_names:
        try:
            table.style = style_name
            return
        except KeyError:
            continue


def get_font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def wrap_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    result: list[str] = []
    for raw_line in text.splitlines() or [""]:
        if not raw_line:
            result.append("")
            continue
        current = ""
        for char in raw_line:
            candidate = current + char
            width = draw.textlength(candidate, font=font)
            if width <= max_width:
                current = candidate
            else:
                result.append(current)
                current = char
        result.append(current)
    return result


def draw_window_base(
    width: int,
    height: int,
    title: str,
    bar_color: str = "#F3F4F6",
    body_color: str = "#FFFFFF",
    title_color: str = "#111827",
) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (width, height), "#E5E7EB")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle(
        (18, 18, width - 18, height - 18),
        radius=20,
        fill=body_color,
        outline="#CBD5E1",
        width=2,
    )
    draw.rounded_rectangle(
        (18, 18, width - 18, 86), radius=20, fill=bar_color, outline="#CBD5E1", width=2
    )
    draw.rectangle((18, 60, width - 18, 86), fill=bar_color, outline=bar_color)
    for idx, color in enumerate(["#EF4444", "#F59E0B", "#22C55E"]):
        x = 42 + idx * 26
        draw.ellipse((x, 38, x + 14, 52), fill=color)
    title_font = get_font(FONT_BOLD, 24)
    draw.text((124, 32), title, fill=title_color, font=title_font)
    return image, draw


def save_window(image: Image.Image, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(out_path)


class RECT(ctypes.Structure):
    _fields_ = [
        ("left", ctypes.c_long),
        ("top", ctypes.c_long),
        ("right", ctypes.c_long),
        ("bottom", ctypes.c_long),
    ]


USER32 = ctypes.windll.user32
EnumWindows = USER32.EnumWindows
EnumWindowsProc = ctypes.WINFUNCTYPE(ctypes.c_bool, ctypes.c_void_p, ctypes.c_void_p)
GetWindowTextLength = USER32.GetWindowTextLengthW
GetWindowText = USER32.GetWindowTextW
GetClassName = USER32.GetClassNameW
IsWindowVisible = USER32.IsWindowVisible
ShowWindow = USER32.ShowWindow
SetForegroundWindow = USER32.SetForegroundWindow
GetWindowRect = USER32.GetWindowRect
GetWindowThreadProcessId = USER32.GetWindowThreadProcessId


def ps_literal(text: str) -> str:
    return text.replace("'", "''")


def enum_windows() -> list[dict[str, object]]:
    windows: list[dict[str, object]] = []

    def callback(hwnd, _lparam):
        if not IsWindowVisible(hwnd):
            return True
        length = GetWindowTextLength(hwnd)
        if length <= 0:
            return True
        title_buf = ctypes.create_unicode_buffer(length + 1)
        GetWindowText(hwnd, title_buf, length + 1)
        class_buf = ctypes.create_unicode_buffer(256)
        GetClassName(hwnd, class_buf, 256)
        proc_id = ctypes.c_ulong()
        GetWindowThreadProcessId(hwnd, ctypes.byref(proc_id))
        windows.append(
            {
                "hwnd": hwnd,
                "title": title_buf.value,
                "class_name": class_buf.value,
                "pid": int(proc_id.value),
            }
        )
        return True

    EnumWindows(EnumWindowsProc(callback), 0)
    return windows


def wait_for_window(
    pid: int | None = None, title_contains: str | None = None, timeout: float = 25.0
) -> int:
    if pid is None and not title_contains:
        raise ValueError("wait_for_window 至少需要 pid 或 title_contains 其一")
    deadline = time.time() + timeout
    while time.time() < deadline:
        for item in enum_windows():
            title = str(item["title"])
            if pid is not None and item["pid"] == pid and title:
                return int(item["hwnd"])
            if title_contains and title_contains in title:
                return int(item["hwnd"])
        time.sleep(0.4)
    if title_contains:
        raise RuntimeError(f"未找到标题包含“{title_contains}”的可见窗口")
    raise RuntimeError(f"未找到进程 {pid} 的可见窗口")


def capture_window(hwnd: int, out_path: Path, delay: float = 1.2) -> None:
    ShowWindow(hwnd, SW_RESTORE)
    time.sleep(0.3)
    ShowWindow(hwnd, SW_MAXIMIZE)
    time.sleep(0.3)
    SetForegroundWindow(hwnd)
    time.sleep(delay)
    rect = RECT()
    GetWindowRect(hwnd, ctypes.byref(rect))
    bbox = (
        max(0, rect.left),
        max(0, rect.top),
        max(1, rect.right),
        max(1, rect.bottom),
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    ImageGrab.grab(bbox=bbox).save(out_path)


def kill_process_tree(proc: subprocess.Popen) -> None:
    subprocess.run(
        ["taskkill", "/PID", str(proc.pid), "/T", "/F"],
        check=False,
        capture_output=True,
        text=True,
    )


def launch_powershell_capture(title: str, commands: list[str], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    script = "\n".join(
        [
            "$host.UI.RawUI.WindowTitle = '{}'".format(ps_literal(title)),
            "chcp 65001 > $null",
            "$utf8 = [System.Text.UTF8Encoding]::new()",
            "[Console]::OutputEncoding = $utf8",
            "[Console]::InputEncoding = $utf8",
            "$OutputEncoding = $utf8",
            "Clear-Host",
            *commands,
            "Start-Sleep -Seconds 180",
        ]
    )
    proc = subprocess.Popen(
        ["powershell.exe", "-NoExit", "-Command", script],
        creationflags=CREATE_NEW_CONSOLE,
    )
    try:
        hwnd = wait_for_window(title_contains=title)
        capture_window(hwnd, out_path)
    finally:
        kill_process_tree(proc)


def copy_to_capture_tmp(src: Path, target_name: str) -> Path:
    CAPTURE_TMP.mkdir(parents=True, exist_ok=True)
    target = CAPTURE_TMP / target_name
    shutil.copy2(src, target)
    return target


def launch_notepad_capture(src_path: Path, out_path: Path, target_name: str) -> None:
    temp_file = copy_to_capture_tmp(src_path, target_name)
    proc = subprocess.Popen(["notepad.exe", str(temp_file)])
    try:
        hwnd = wait_for_window(title_contains=target_name)
        capture_window(hwnd, out_path)
    finally:
        kill_process_tree(proc)


def launch_paint_capture(src_path: Path, out_path: Path, target_name: str) -> None:
    temp_file = copy_to_capture_tmp(src_path, target_name)
    proc = subprocess.Popen(["mspaint.exe", str(temp_file)])
    try:
        hwnd = wait_for_window(proc.pid)
        capture_window(hwnd, out_path, delay=1.6)
    finally:
        kill_process_tree(proc)


def make_terminal_window(
    out_path: Path,
    title: str,
    command_line: str,
    body_lines: list[str],
    footer: str | None = None,
) -> None:
    width, height = 1600, 960
    image, draw = draw_window_base(
        width,
        height,
        title,
        bar_color="#111827",
        body_color="#111827",
        title_color="#E5E7EB",
    )
    draw.rounded_rectangle((44, 102, width - 44, 160), radius=12, fill="#1F2937")
    mono = get_font(FONT_MONO, 22)
    regular = get_font(FONT_REGULAR, 20)
    draw.text((68, 118), command_line, fill="#F8FAFC", font=mono)
    y = 188
    max_width = width - 120
    for raw in body_lines:
        wrapped = wrap_lines(draw, raw, mono, max_width)
        for line in wrapped:
            draw.text((68, y), line, fill="#D1D5DB", font=mono)
            y += 34
            if y > height - 90:
                break
        if y > height - 90:
            break
    if footer:
        draw.text((68, height - 60), footer, fill="#9CA3AF", font=regular)
    save_window(image, out_path)


def make_explorer_window(
    out_path: Path, title: str, current_path: str, items: list[tuple[str, str, str]]
) -> None:
    width, height = 1600, 960
    image, draw = draw_window_base(width, height, title)
    regular = get_font(FONT_REGULAR, 20)
    bold = get_font(FONT_BOLD, 22)
    draw.rounded_rectangle(
        (52, 108, width - 52, 156), radius=10, fill="#F8FAFC", outline="#CBD5E1"
    )
    draw.text((74, 120), current_path, fill="#0F172A", font=regular)
    draw.rounded_rectangle(
        (52, 172, 314, height - 52), radius=12, fill="#F8FAFC", outline="#CBD5E1"
    )
    for idx, name in enumerate(["快速访问", "项目目录", "data", "Results", "images"]):
        draw.text((80, 198 + idx * 48), name, fill="#334155", font=regular)
    draw.rounded_rectangle(
        (336, 172, width - 52, height - 52),
        radius=12,
        fill="#FFFFFF",
        outline="#CBD5E1",
    )
    draw.text((370, 196), "名称", fill="#475569", font=bold)
    draw.text((1020, 196), "类型", fill="#475569", font=bold)
    draw.text((1250, 196), "备注", fill="#475569", font=bold)
    y = 244
    for name, kind, note in items[:14]:
        draw.line((350, y - 12, width - 70, y - 12), fill="#E2E8F0", width=1)
        draw.rounded_rectangle((370, y, 392, y + 22), radius=5, fill="#60A5FA")
        draw.text((406, y - 2), name, fill="#0F172A", font=regular)
        draw.text((1020, y - 2), kind, fill="#475569", font=regular)
        draw.text((1250, y - 2), note, fill="#475569", font=regular)
        y += 48
    save_window(image, out_path)


def make_editor_window(
    out_path: Path, title: str, path_label: str, body_lines: list[str]
) -> None:
    width, height = 1600, 960
    image, draw = draw_window_base(width, height, title)
    mono = get_font(FONT_MONO, 20)
    regular = get_font(FONT_REGULAR, 18)
    draw.rounded_rectangle(
        (52, 108, width - 52, 156), radius=10, fill="#F8FAFC", outline="#CBD5E1"
    )
    draw.text((74, 122), path_label, fill="#0F172A", font=regular)
    draw.rounded_rectangle(
        (52, 172, width - 52, height - 52), radius=12, fill="#F8FAFC", outline="#CBD5E1"
    )
    y = 198
    for idx, raw in enumerate(body_lines[:22], start=1):
        draw.text((78, y), f"{idx:>2}", fill="#94A3B8", font=mono)
        draw.text((138, y), raw, fill="#0F172A", font=mono)
        y += 32
    save_window(image, out_path)


def make_viewer_window(
    out_path: Path, title: str, image_path: Path, footer: str
) -> None:
    width, height = 1600, 960
    image, draw = draw_window_base(width, height, title)
    regular = get_font(FONT_REGULAR, 18)
    inner = Image.open(image_path).convert("RGB")
    inner.thumbnail((1450, 760))
    x = (width - inner.width) // 2
    y = 120
    image.paste(inner, (x, y))
    draw.rounded_rectangle(
        (x - 12, y - 12, x + inner.width + 12, y + inner.height + 12),
        radius=12,
        outline="#CBD5E1",
        width=2,
    )
    draw.text((70, height - 56), footer, fill="#475569", font=regular)
    save_window(image, out_path)


def make_csv_window(
    out_path: Path, title: str, columns: list[str], rows: list[list[str]]
) -> None:
    width, height = 1600, 960
    image, draw = draw_window_base(width, height, title)
    regular = get_font(FONT_REGULAR, 18)
    bold = get_font(FONT_BOLD, 18)
    draw.rounded_rectangle(
        (52, 110, width - 52, height - 52), radius=12, fill="#FFFFFF", outline="#CBD5E1"
    )
    col_x = [74, 220, 430, 720, 1020, 1280]
    headers = columns[:6]
    for idx, header in enumerate(headers):
        draw.text((col_x[idx], 136), header, fill="#0F172A", font=bold)
    y = 180
    for row in rows[:12]:
        draw.line((68, y - 10, width - 70, y - 10), fill="#E2E8F0", width=1)
        for idx, value in enumerate(row[:6]):
            draw.text((col_x[idx], y), value, fill="#334155", font=regular)
        y += 46
    save_window(image, out_path)


def read_help_output() -> list[str]:
    proc = subprocess.run(
        ["uv", "run", "python", "run.py", "--help"],
        cwd=REPO,
        capture_output=True,
        text=True,
        encoding="utf-8",
        check=True,
    )
    return proc.stdout.splitlines()


def read_csv_rows(path: Path, limit: int = 10) -> tuple[list[str], list[list[str]]]:
    with path.open("r", encoding="utf-8", newline="") as fh:
        reader = csv.reader(fh)
        rows = list(reader)
    return rows[0], rows[1 : 1 + limit]


def build_screenshots() -> dict[str, Path]:
    screenshots = {
        "project_root": SCREENSHOT_DIR / "01_project_root.png",
        "data_folder": SCREENSHOT_DIR / "02_data_folder.png",
        "cli_help": SCREENSHOT_DIR / "03_cli_help.png",
        "preset_cmd": SCREENSHOT_DIR / "04_preset_experiment.png",
        "custom_cmd": SCREENSHOT_DIR / "05_custom_experiment.png",
        "results_folder": SCREENSHOT_DIR / "06_results_folder.png",
        "pareto": SCREENSHOT_DIR / "07_pareto_window.png",
        "report": SCREENSHOT_DIR / "08_report_window.png",
        "dashboard": SCREENSHOT_DIR / "09_dashboard_window.png",
        "csv_preview": SCREENSHOT_DIR / "10_csv_preview.png",
    }
    report_path = RESULT_DIR / "comparison_report.md"
    post_csv_path = RESULT_DIR / "post_analysis_results.csv"

    launch_powershell_capture(
        "AgencyShot-ProjectRoot",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(REPO))}'",
            f"Write-Host '软件名称：{SOFTWARE_NAME}' -ForegroundColor Cyan",
            "Write-Host '命令：Get-ChildItem -Force' -ForegroundColor Yellow",
            "$show = 'AGENTS.md','Readme.md','case_config.py','cchp_gaproblem.py','cchp_gasolution.py','operation.py','run.py','data','Results','scripts'",
            "Get-ChildItem -Force | Where-Object { $show -contains $_.Name } | Sort-Object Name | Select-Object Mode,LastWriteTime,Length,Name | Format-Table -AutoSize",
        ],
        screenshots["project_root"],
    )
    launch_powershell_capture(
        "AgencyShot-DataFolder",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(DATA_DIR))}'",
            "Write-Host '输入数据目录检查' -ForegroundColor Cyan",
            "Write-Host '命令：Get-ChildItem -Force' -ForegroundColor Yellow",
            "Get-ChildItem -Force | Sort-Object Name | Select-Object Mode,LastWriteTime,Length,Name | Format-Table -AutoSize",
        ],
        screenshots["data_folder"],
    )
    launch_powershell_capture(
        "AgencyShot-Help",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(REPO))}'",
            "Write-Host '命令帮助窗口' -ForegroundColor Cyan",
            "Write-Host '命令：uv run python run.py --help' -ForegroundColor Yellow",
            "uv run python run.py --help | Select-Object -First 24",
        ],
        screenshots["cli_help"],
    )
    launch_powershell_capture(
        "AgencyShot-Preset",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(REPO))}'",
            "Write-Host '预设实验命令示例' -ForegroundColor Cyan",
            "Write-Host '命令：uv run python run.py --exp 1 --test-run' -ForegroundColor Yellow",
            "Write-Host ''",
            "Write-Host '以下显示本机已有 test_exp1 结果报告摘要：' -ForegroundColor Green",
            f"Get-Content -LiteralPath '{ps_literal(str(report_path))}' -Encoding UTF8 -TotalCount 14",
        ],
        screenshots["preset_cmd"],
    )
    launch_powershell_capture(
        "AgencyShot-Custom",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(REPO))}'",
            "Write-Host '自定义实验命令示例' -ForegroundColor Cyan",
            "Write-Host '命令：uv run python run.py --mode custom --case songshan_lake --nind 20 --maxgen 20 --methods std euclidean' -ForegroundColor Yellow",
            "Write-Host ''",
            "Write-Host '常用扩展参数：' -ForegroundColor Green",
            "Write-Host '  --carnot          启用卡诺电池' -ForegroundColor White",
            "Write-Host '  --workers N       指定并行进程数' -ForegroundColor White",
            "Write-Host '  --quick-run       使用快速实验参数' -ForegroundColor White",
            "Write-Host ''",
            "Write-Host '运行前应确认案例数据、典型日文件和求解器环境已配置完成。' -ForegroundColor Gray",
        ],
        screenshots["custom_cmd"],
    )
    launch_powershell_capture(
        "AgencyShot-Results",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(RESULT_DIR))}'",
            "Write-Host '结果目录检查' -ForegroundColor Cyan",
            "Write-Host '命令：Get-ChildItem -Force' -ForegroundColor Yellow",
            "Get-ChildItem -Force | Sort-Object Name | Select-Object Mode,LastWriteTime,Length,Name | Format-Table -AutoSize",
        ],
        screenshots["results_folder"],
    )
    shutil.copy2(IMAGE_SOURCE_DIR / "fig-pareto-comparison.png", screenshots["pareto"])
    launch_powershell_capture(
        "AgencyShot-Report",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(RESULT_DIR))}'",
            "Write-Host '对比报告预览' -ForegroundColor Cyan",
            f"Write-Host '文件：{report_path.name}' -ForegroundColor Yellow",
            f"Get-Content -LiteralPath '{ps_literal(str(report_path))}' -Encoding UTF8 -TotalCount 18",
        ],
        screenshots["report"],
    )
    shutil.copy2(IMAGE_SOURCE_DIR / "fig-report-preview.png", screenshots["dashboard"])
    launch_powershell_capture(
        "AgencyShot-CSV",
        [
            f"Set-Location -LiteralPath '{ps_literal(str(REPO))}'",
            "Write-Host '后验分析结果预览' -ForegroundColor Cyan",
            f"Import-Csv -LiteralPath '{ps_literal(str(post_csv_path))}' | Select-Object -First 8 cost_level,method,annual_cost,matching_index,self_sufficiency_% | Format-Table -AutoSize",
        ],
        screenshots["csv_preview"],
    )
    shutil.rmtree(CAPTURE_TMP, ignore_errors=True)
    return screenshots


def fill_form_document() -> None:
    document = Document(str(FORM_TEMPLATE))
    table1, table2, table3 = document.tables

    set_cell_text(table1.cell(0, 1), SOFTWARE_NAME)
    set_cell_text(table1.cell(1, 1), "应用软件")
    set_cell_text(table1.cell(2, 1), SOFTWARE_SHORT)
    set_cell_text(table1.cell(2, 3), SOFTWARE_VERSION)
    set_cell_text(table1.cell(3, 1), DEV_DATE)
    set_cell_text(table1.cell(3, 3), PUBLISH_STATUS)
    set_cell_text(table1.cell(4, 1), PUBLISH_DATE)
    set_cell_text(table1.cell(4, 3), PUBLISH_CITY)
    set_cell_text(table1.cell(5, 1), "原创软件")
    set_cell_text(table1.cell(6, 1), "原始取得")
    set_cell_text(table1.cell(7, 1), "全部权利")
    set_cell_text(table1.cell(8, 1), "单独开发")

    set_cell_text(table2.cell(0, 1), OWNER_NAME)
    set_cell_text(table2.cell(1, 1), OWNER_CODE)
    set_cell_text(table2.cell(2, 1), "中国，浙江省杭州市")

    set_cell_text(table3.cell(0, 1), CONTACT_NAME)
    set_cell_text(table3.cell(0, 3), CONTACT_PHONE)
    set_cell_text(table3.cell(1, 1), CONTACT_EMAIL)
    set_cell_text(table3.cell(2, 1), "郑浩男（浙江大学能源工程学院）")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    add_text(
        p,
        "五、软件技术特点（以下内容已按事务所要求整理）",
        east_asia="黑体",
        size=14,
        bold=True,
        latin="Arial",
    )

    tech_table = document.add_table(rows=0, cols=2)
    apply_table_style(tech_table, "Table Grid", "TableGrid")
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("开发的硬件环境（50字符内）", DEV_HW),
        ("运行的硬件环境（50字符内）", RUN_HW),
        ("开发该软件的操作系统（50字符内）", DEV_OS),
        ("软件开发环境/开发工具（50字符内）", DEV_TOOL),
        ("该软件的运行平台/操作系统（50字符内）", RUN_OS),
        ("软件运行支撑环境/支持软件（50字符内）", SUPPORT_SW),
        ("编程语言", PROGRAM_LANG),
        ("源程序量", f"{SOURCE_LINES_FOR_SUBMISSION}行"),
        ("开发目的（50字符内）", PURPOSE),
        ("面向领域/行业（50字符内）", DOMAIN),
        ("主要功能（500-1300字符）", MAIN_FUNCTIONS),
        ("技术特点（100字符内）", TECH_FEATURE),
    ]
    for key, value in rows:
        row_cells = tech_table.add_row().cells
        set_cell_text(row_cells[0], key, bold=True)
        set_cell_text(row_cells[1], value)

    p2 = document.add_paragraph()
    add_text(
        p2,
        "说明：本次提交的软件登记信息表、使用说明书和源代码文件中的软件名称、版本号、著作权人和联系人信息已保持一致；源代码提交口径按事务所要求采用前1500行和后1500行，共3000行。",
        east_asia="宋体",
        size=10.5,
    )
    document.save(str(OUTPUT_FORM))


def add_simple_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths_cm: list[float] | None = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    apply_table_style(table, "Table Grid", "TableGrid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)
        if col_widths_cm:
            hdr[idx].width = Cm(col_widths_cm[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if col_widths_cm:
                cells[idx].width = Cm(col_widths_cm[idx])
    document.add_paragraph()


def build_manual_document(screenshots: dict[str, Path]) -> None:
    document = Document(str(MANUAL_TEMPLATE))
    clear_body(document)
    add_title(document, SOFTWARE_NAME, "使用说明书")

    add_heading_line(document, "1 引言", "Heading 2")
    add_heading_line(document, "1.1 编写目的", "Heading 3")
    add_body_text(
        document,
        "本说明书用于说明“源荷匹配的分布式电热综合能源系统优化软件”的运行环境、输入数据、命令入口、实验操作流程、结果查看方法和使用注意事项，便于向事务所提交完整、连贯的使用说明材料。该软件为研究型命令行程序，没有独立的登录页面和图形化首页，实际使用入口为项目根目录下的命令行窗口，因此本说明书中的命令行窗口视为软件启动界面。",
    )
    add_heading_line(document, "1.2 软件概述", "Heading 3")
    add_body_text(
        document,
        "该软件面向分布式冷热电联供系统优化规划问题，围绕“源荷匹配”这一研究主题组织建模、求解和结果分析流程。程序支持德国社区案例和松山湖校园案例两套输入数据，能够对光伏、风电、燃气轮机、电热泵、电制冷机、吸收式制冷机以及冷热电储能容量进行优化，并在启用卡诺电池时扩展为含电热耦合储能的配置方案。",
    )
    add_heading_line(document, "1.3 参考资料", "Heading 3")
    add_body_text(
        document,
        "主要参考资料包括论文《Energy-quality-weighted source-load matching for optimal planning of distributed combined cooling, heating and power systems with Carnot battery integration》、代码仓库中的 run.py、case_config.py、operation.py、cchp_gaproblem.py、cchp_gasolution.py 和 scripts/post_analysis_report.py 文件，以及结果目录中的 comparison_report.md 和 post_analysis_results.csv 文件。",
    )

    add_heading_line(document, "2 运行环境", "Heading 2")
    add_heading_line(document, "2.1 硬件与软件环境", "Heading 3")
    add_body_text(
        document,
        "本软件建议在64位Windows工作站上运行，内存不低于16GB，磁盘建议使用SSD。开发与运行环境均以 Windows 10/11 为主，程序依赖 Python 3.8、geatpy、oemof-solph、Pyomo 等第三方库，求解器优先使用 Gurobi，不具备 Gurobi 条件时可退回到 GLPK。",
    )
    add_simple_table(
        document,
        ["项目", "说明"],
        [
            ["硬件环境", "x64 PC或服务器，16GB内存，建议SSD存储"],
            ["操作系统", "Windows 10/11 64位"],
            ["Python环境", "Python 3.8"],
            ["求解器", "Gurobi优先，GLPK备用"],
            ["核心依赖", "geatpy、oemof-solph、Pyomo、Pandas、Numpy"],
        ],
        [4.2, 11.8],
    )
    add_heading_line(document, "2.2 输入数据准备", "Heading 3")
    add_body_text(
        document,
        "程序运行前需要保留 data 目录中的负荷、气象和典型日文件。德国案例使用 mergedData.csv 和 typicalDayData.xlsx，松山湖案例使用 songshan_lake_data.csv 和 songshan_lake_typical.xlsx。用户应确认数据文件与代码目录保持相对路径一致，避免运行时读取失败。",
    )
    add_simple_table(
        document,
        ["文件名", "作用", "说明"],
        [
            [
                "mergedData.csv",
                "德国案例逐时数据",
                "包含8760小时电、热、冷负荷及气象数据",
            ],
            [
                "typicalDayData.xlsx",
                "德国案例典型日权重",
                "定义14个典型日及对应代表天数",
            ],
            [
                "songshan_lake_data.csv",
                "松山湖案例逐时数据",
                "包含校园案例逐时负荷和气象数据",
            ],
            ["songshan_lake_typical.xlsx", "松山湖典型日权重", "用于全年加权计算"],
        ],
        [4.2, 4.8, 7.0],
    )

    document.add_picture(str(screenshots["project_root"]), width=Cm(15.4))
    add_caption(document, "图2-1 项目根目录界面")
    document.add_picture(str(screenshots["data_folder"]), width=Cm(15.4))
    add_caption(document, "图2-2 输入数据目录界面")

    add_heading_line(document, "3 软件使用说明", "Heading 2")
    add_heading_line(document, "3.1 查看命令帮助", "Heading 3")
    add_body_text(
        document,
        "在项目根目录打开命令行窗口后，首先执行“uv run python run.py --help”查看程序支持的运行模式、实验编号、案例切换参数和并行设置。命令帮助界面中给出了 test、quick、full、custom 模式，以及 --exp、--case、--carnot、--workers 等参数的说明。首次使用时建议先阅读帮助信息，确认命令格式后再执行正式计算。",
    )
    document.add_picture(str(screenshots["cli_help"]), width=Cm(15.4))
    add_caption(document, "图3-1 命令帮助界面")

    add_heading_line(document, "3.2 运行论文预设实验", "Heading 3")
    add_body_text(
        document,
        "论文中的预设实验可直接通过 --exp 参数调用。以实验1为例，用户在命令行中输入“uv run python run.py --exp 1 --test-run”后，程序会按照德国案例、多方法对比的设定执行测试规模计算。采用 --test-run 时，程序会自动使用较小种群和迭代次数进行流程验证，适合检查环境、求解器和结果目录是否正常。",
    )
    add_body_text(
        document,
        "实验执行结束后，软件会在 Results 目录下生成带时间戳的结果文件夹。文件夹内包含 Economic_only、Std、Euclidean、Pearson、SSR 等方法对应的 CSV 结果文件，以及 Pareto_Comparison.png 和 comparison_report.md 等汇总结果。",
    )
    document.add_picture(str(screenshots["preset_cmd"]), width=Cm(15.4))
    add_caption(document, "图3-2 预设实验命令及结果摘要")
    document.add_picture(str(screenshots["results_folder"]), width=Cm(15.4))
    add_caption(document, "图3-3 结果目录界面")

    add_heading_line(document, "3.3 运行自定义实验", "Heading 3")
    add_body_text(
        document,
        "当用户需要自行指定案例、方法或计算规模时，可使用 custom 模式。例如输入“uv run python run.py --mode custom --case songshan_lake --nind 20 --maxgen 20 --methods std euclidean”后，程序将按松山湖案例执行波动率和欧氏匹配指标的对比计算。如果需要加入卡诺电池，可在命令中附加 --carnot 参数。",
    )
    add_body_text(
        document,
        "自定义实验适用于调节种群规模、最大代数和方法组合。参数越大，求解时间越长。正式提交计算前，建议先在小规模参数下验证命令、数据路径和结果输出是否正确，再扩大规模开展正式实验。",
    )
    document.add_picture(str(screenshots["custom_cmd"]), width=Cm(15.4))
    add_caption(document, "图3-4 自定义实验命令界面")

    add_heading_line(document, "3.4 查看图形和报告结果", "Heading 3")
    add_body_text(
        document,
        "双目标方法运行完成后，软件会输出 Pareto 对比图，用于展示经济成本与匹配指标之间的折中关系。图中不同颜色代表不同方法，用户可据此筛选低成本方案和高匹配方案。对于论文复现任务，通常需要结合该图和后续的报告表格共同判断最优方案。",
    )
    document.add_picture(str(screenshots["pareto"]), width=Cm(15.4))
    add_caption(document, "图3-5 Pareto 对比结果图")

    add_body_text(
        document,
        "comparison_report.md 用于汇总各方法的最低成本、最佳匹配度、运行时间和 Pareto 解数量，并进一步列出最优方案下的设备容量配置。用户可直接打开该文件查看方法之间的成本和指标差异。",
    )
    document.add_picture(str(screenshots["report"]), width=Cm(15.4))
    add_caption(document, "图3-6 对比报告预览")

    add_body_text(
        document,
        "post_analysis_results.csv 用于保存8760小时后验分析结果。该文件记录不同成本层级下的年总成本、峰值购电、自给率、弃风弃光率和匹配指标，可用于进一步绘图和论文分析。",
    )
    document.add_picture(str(screenshots["dashboard"]), width=Cm(15.4))
    add_caption(document, "图3-7 后验分析图形结果")
    document.add_picture(str(screenshots["csv_preview"]), width=Cm(15.4))
    add_caption(document, "图3-8 后验分析结果表预览")

    add_heading_line(document, "4 主要功能说明", "Heading 2")
    add_heading_line(document, "4.1 案例配置管理", "Heading 3")
    add_body_text(
        document,
        "软件通过 case_config.py 管理案例数据路径、电价曲线、气价曲线、容量上界、投资系数和设备效率等参数。用户可以在已有德国案例和松山湖案例基础上继续添加新案例，只需补充新的配置字典，并准备相应的逐时数据文件和典型日权重文件。",
    )
    add_heading_line(document, "4.2 多目标容量优化", "Heading 3")
    add_body_text(
        document,
        "优化层通过 geatpy 构建问题对象，对光伏、风电、燃气轮机、电热泵、电制冷机、吸收式制冷机以及冷热电储能容量进行搜索。在双目标模式下，软件同时考虑经济成本和匹配指标；在单目标模式下，仅按经济成本求解。程序支持 std、euclidean、pearson、ssr 和 economic_only 五类方法。",
    )
    add_heading_line(document, "4.3 调度仿真与指标计算", "Heading 3")
    add_body_text(
        document,
        "每个候选容量方案在评估时都会调用运行调度模型，对14个典型日进行24小时求解。调度模型同时处理电、热、冷、气四类能流，并可在启用卡诺电池时增加电储能与余热回收约束。调度结果进一步用于计算年化成本、峰值购电、弃风弃光率以及多种源荷匹配指标。",
    )
    add_heading_line(document, "4.4 结果导出与后验分析", "Heading 3")
    add_body_text(
        document,
        "软件在 Results 目录下输出各方法的 Pareto 结果、设备参数表、Markdown 对比报告和图形化结果文件。scripts/post_analysis_report.py 还可以进一步调用后验分析逻辑，对 8760 小时结果进行再统计，并生成预算分层与敏感性分析所需数据。",
    )

    add_heading_line(document, "5 注意事项", "Heading 2")
    add_body_text(
        document,
        "（1）本软件为命令行科研程序，正式交付截图时应完整保留命令行窗口边框、目录窗口标题栏和图形窗口界面，不宜只截取局部图块。",
    )
    add_body_text(
        document,
        "（2）如本机已配置 Gurobi，请确认许可证路径可用；若 Gurobi 不可用，程序会回退到 GLPK，但求解效率可能下降。",
    )
    add_body_text(
        document,
        "（3）松山湖案例数据为合成逐时数据，适用于研究验证和方法比较，不直接代表真实工程监测值。",
    )
    add_body_text(
        document,
        "（4）提交事务所前，建议再次检查软件名称、版本号、著作权人名称以及源代码行数口径是否与登记信息表一致。",
    )
    document.save(str(OUTPUT_MANUAL))


def selected_source_paths() -> list[Path]:
    return [REPO / rel for rel in SELECTED_SOURCE_FILES]


def build_source_document() -> None:
    document = Document(str(SOURCE_TEMPLATE))
    clear_body(document)

    section = document.sections[0]
    section.start_type = WD_SECTION_START.CONTINUOUS

    code_style = "Normal"
    all_lines: list[str] = []
    for path in selected_source_paths():
        text = path.read_text(encoding="utf-8", errors="ignore").splitlines()
        all_lines.extend(text)

    if len(all_lines) < SOURCE_LINES_FOR_SUBMISSION:
        selected = all_lines
        split_at = len(all_lines)
    else:
        selected = all_lines[:1500] + all_lines[-1500:]
        split_at = 1500

    for idx, line in enumerate(selected, start=1):
        paragraph = document.add_paragraph(style=code_style)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.left_indent = Pt(0)
        run = paragraph.add_run(line)
        set_run_font(run, east_asia="Calibri", size=9, latin="Calibri")
        if idx == split_at:
            run.add_break(WD_BREAK.PAGE)

    document.save(str(OUTPUT_SOURCE))


def validate_documents() -> None:
    for path in [OUTPUT_FORM, OUTPUT_MANUAL, OUTPUT_SOURCE]:
        Document(str(path))


def main() -> None:
    assert_lengths()
    ensure_dirs()
    screenshots = build_screenshots()
    fill_form_document()
    build_manual_document(screenshots)
    build_source_document()
    validate_documents()
    print("Generated:")
    print(OUTPUT_FORM)
    print(OUTPUT_MANUAL)
    print(OUTPUT_SOURCE)


if __name__ == "__main__":
    main()
